# -*- coding: utf-8 -*-
"""Создание Word-документа о Nissan Skyline."""

import io
import os
import time
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(OUTPUT_DIR, "images")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Nissan_Skyline.docx")

# Реальные фото с Wikimedia Commons (thumb-версии 1280px)
IMAGES = {
    "r34": "https://upload.wikimedia.org/wikipedia/commons/thumb/0/06/Nissan_Skyline_GT-R_R34_V_Spec_II.jpg/1280px-Nissan_Skyline_GT-R_R34_V_Spec_II.jpg",
    "r32": "https://upload.wikimedia.org/wikipedia/commons/thumb/0/00/Nissan_SKYLINE_GT-R_%28E-BNR32%29_front.JPG/1280px-Nissan_SKYLINE_GT-R_%28E-BNR32%29_front.JPG",
    "r33": "https://upload.wikimedia.org/wikipedia/commons/thumb/8/88/1996_Nissan_Skyline_GT-R_R33_V-Spec.jpg/1280px-1996_Nissan_Skyline_GT-R_R33_V-Spec.jpg",
    "v35": "https://upload.wikimedia.org/wikipedia/commons/thumb/8/8a/Japanese_NISSAN_Skyline_V35.jpg/1280px-Japanese_NISSAN_Skyline_V35.jpg",
    "r35_gtr": "https://upload.wikimedia.org/wikipedia/commons/thumb/e/ef/2009-2010_Nissan_GT-R_%28R35%29_coupe_01.jpg/1280px-2009-2010_Nissan_GT-R_%28R35%29_coupe_01.jpg",
}


def download_image(url, filename, key=None, force=False):
    """Скачать изображение с Wikimedia Commons."""
    os.makedirs(IMAGES_DIR, exist_ok=True)
    path = os.path.join(IMAGES_DIR, filename)
    if os.path.exists(path) and os.path.getsize(path) > 8000 and not force:
        return path
    if force and os.path.exists(path):
        os.remove(path)
    headers = {"User-Agent": "SkylineDocBot/1.0 (educational)"}
    try:
        time.sleep(2)
        resp = requests.get(url, headers=headers, timeout=60)
        resp.raise_for_status()
        if len(resp.content) < 8000:
            raise ValueError("file too small")
        with open(path, "wb") as f:
            f.write(resp.content)
        return path
    except Exception as e:
        print(f"  Download failed {filename}: {e}")
        if os.path.exists(path) and os.path.getsize(path) > 8000:
            return path
        return None


def set_cell_shading(cell, color_hex):
    """Заливка ячейки таблицы."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc, color="C41E3A"):
    """Декоративная линия."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_styles(doc):
    """Настройка стилей документа."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size, color in [(1, 22, "1A1A2E"), (2, 16, "C41E3A"), (3, 13, "16213E")]:
        name = f"Heading {level}"
        h = doc.styles[name]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor.from_string(color)
        h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    """Титульная страница."""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NISSAN SKYLINE")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(26, 26, 46)
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Легенда японского автопрома")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(196, 30, 58)
    run.italic = True
    run.font.name = "Calibri"

    add_horizontal_line(doc)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("История · Поколения · GT-R · Культурное наследие")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def add_image(doc, path, width=Inches(5.5), caption=None):
    """Вставить изображение с подписью."""
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(120, 120, 120)
        doc.add_paragraph()


def add_highlight_box(doc, text):
    """Выделенный блок с информацией."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F5F0F0")
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(26, 26, 46)
    doc.add_paragraph()


def add_specs_table(doc):
    """Таблица характеристик GT-R R34."""
    doc.add_heading("Сравнение поколений Skyline GT-R", level=2)

    headers = ["Поколение", "Годы", "Двигатель", "Мощность", "Привод"]
    rows_data = [
        ["KPGC10 Hakosuka", "1969–1972", "2.0 л S20 I6", "160 л.с.", "Задний"],
        ["R32 GT-R", "1989–1994", "2.6 л RB26DETT", "280 л.с.", "AWD ATTESA"],
        ["R33 GT-R", "1995–1998", "2.6 л RB26DETT", "280 л.с.", "AWD ATTESA"],
        ["R34 GT-R", "1999–2002", "2.6 л RB26DETT", "280 л.с.", "AWD ATTESA"],
        ["R35 GT-R", "2007–н.в.", "3.8 л VR38DETT", "570+ л.с.", "AWD ATTESA"],
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1A1A2E")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 0:
                set_cell_shading(cell, "F8F8FA")
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph()


def build_document(image_paths):
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_title_page(doc)

    # --- Введение ---
    doc.add_heading("Введение", level=1)
    add_horizontal_line(doc)

    doc.add_paragraph(
        "Nissan Skyline — одна из самых узнаваемых и почитаемых серий автомобилей "
        "в истории японского автопрома. С 1957 года эта модель прошла путь от скромного "
        "седана среднего класса до иконы мирового автоспорта и уличных гонок. "
        "Особую славу Skyline принесла линейка GT-R — автомобили, которые навсегда "
        "изменили представление о возможностях японских спортивных машин."
    )

    if image_paths.get("r34"):
        add_image(doc, image_paths["r34"],
                  caption="Nissan Skyline GT-R R34 V-Spec II — культовое поколение")

    add_highlight_box(
        doc,
        "💡 Интересный факт: обозначение «GT-R» расшифровывается как "
        "«Gran Turismo Racing» — гоночный гран-туризмо. Эта маркировка появилась "
        "ещё в 1969 году на модели Skyline 2000GT-R (кодовое имя Hakosuka)."
    )

    # --- История ---
    doc.add_heading("История создания", level=1)
    add_horizontal_line(doc)

    doc.add_heading("Ранние годы (1957–1968)", level=2)
    doc.add_paragraph(
        "Первый Nissan Skyline (модель ALSI-1) был представлен в 1957 году компанией "
        "Prince Motor Company, которая позже влилась в Nissan. Изначально Skyline "
        "позиционировался как премиальный седан для японского рынка. В 1964 году "
        "Prince Skyline GT завоевал второе место в Гран-при Японии, положив начало "
        "спортивной репутации марки."
    )

    doc.add_heading("Рождение GT-R (1969–1972)", level=2)
    doc.add_paragraph(
        "В феврале 1969 года дебютировал Skyline 2000GT-R (PGC10) — первый настоящий GT-R. "
        "Оснащённый рядным шестицилиндровым двигателем S20 мощностью 160 л.с., "
        "он доминировал на гоночных трассах Японии, одержав 52 победы за два года. "
        "Эту модель прозвали «Hakosuka» (箱作) — «коробчатый Skyline» за характерный "
        "квадратный дизайн кузова."
    )

    doc.add_heading("Эра «Godzilla» (1989–2002)", level=2)
    doc.add_paragraph(
        "Возрождение GT-R в 1989 году с моделью R32 стало поворотным моментом. "
        "Британский журнал Wheels прозвал R32 GT-R «Godzilla» — монстром, "
        "который уничтожал конкурентов на треке. Двигатель RB26DETT с двумя турбонаддувами, "
        "интеллектуальный полный привод ATTESA E-TS и система активного дифференциала "
        "Super HICAS сделали R32 непобедимым на гонках Group A."
    )

    if image_paths.get("r32"):
        add_image(doc, image_paths["r32"],
                  caption="Nissan Skyline GT-R R32 (E-BNR32) — «Godzilla»")

    if image_paths.get("r33"):
        add_image(doc, image_paths["r33"],
                  caption="Nissan Skyline GT-R R33 V-Spec (1996)")

    # --- Поколения GT-R ---
    doc.add_heading("Поколения Skyline GT-R", level=1)
    add_horizontal_line(doc)

    generations = [
        ("R32 GT-R (1989–1994)", [
            "Двигатель RB26DETT — 2.6 л, рядный 6-цилиндровый, битурбо",
            "Система полного привода ATTESA E-TS Pro",
            "Активное рулевое управление Super HICAS",
            "29 побед из 29 гонок в японском чемпионате touring car",
            "Лимит мощности 280 л.с. по японской «джентльменской соглашению»",
        ]),
        ("R33 GT-R (1995–1998)", [
            "Усовершенствованная подвеска и увеличенный кузов",
            "Впервые GT-R прошёл круг Nürburgring быстрее 8 минут",
            "Модификация V-Spec с активным дифференциалом LSD",
            "400R — лимитированная версия с 400 л.с.",
        ]),
        ("R34 GT-R (1999–2002)", [
            "Мультифункциональный дисплей на приборной панели",
            "Двигатель RB26DETT с улучшенными турбинами",
            "V-Spec II с карбоновым воздухозаборником на капоте",
            "Nismo Z-Tune — 500 л.с., всего 20 экземпляров",
            "Звезда фильмов «Форсаж 2» и аниме «Initial D»",
        ]),
        ("R35 GT-R (2007–н.в.)", [
            "Отдельная модель (без приставки Skyline в названии)",
            "Двигатель VR38DETT — 3.8 л V6, 570–600 л.с.",
            "Ручная сборка двигателя мастером «Takumi»",
            "Рекорд Nürburgring среди серийных авто своего времени",
            "Постоянные обновления: Nismo, GT-R50, T-Spec",
        ]),
    ]

    for title, points in generations:
        doc.add_heading(title, level=2)
        for point in points:
            p = doc.add_paragraph(point, style="List Bullet")

    if image_paths.get("r35_gtr"):
        add_image(doc, image_paths["r35_gtr"],
                  caption="Nissan GT-R (R35) — современное воплощение легенды")

    # --- Таблица ---
    add_specs_table(doc)

    # --- Skyline без GT-R ---
    doc.add_heading("Skyline вне линейки GT-R", level=1)
    add_horizontal_line(doc)

    doc.add_paragraph(
        "Помимо легендарных GT-R, серия Skyline включала множество других моделей, "
        "популярных на японском рынке. Skyline R34 в обычной комплектации предлагал "
        "двигатели от 2.0-литрового RB20DE до турбированного RB25DET. "
        "В 2001 году появился Skyline V35 (продававшийся как Infiniti G35 за рубежом) — "
        "первый Skyline на платформе FM с задним приводом и V6-двигателем VQ."
    )

    if image_paths.get("v35"):
        add_image(doc, image_paths["v35"],
                  caption="Nissan Skyline V35 — новая эра дизайна")

    doc.add_paragraph(
        "Современный Skyline (V37, с 2014 года) базируется на той же платформе, "
        "что и Infiniti Q50, и оснащается турбированными четырёхцилиндровыми "
        "или гибридными силовыми установками. GT-R же стал самостоятельной моделью, "
        "сохранив дух оригинала."
    )

    # --- Культура ---
    doc.add_heading("Культурное наследие", level=1)
    add_horizontal_line(doc)

    doc.add_paragraph(
        "Nissan Skyline GT-R — один из самых культовых автомобилей в мировой "
        "поп-культуре. Его появление в кинофраншизе «Форсаж» (Fast & Furious), "
        "аниме и манге «Initial D», видеоиграх Gran Turismo и Need for Speed "
        "сделало Skyline мечтой целого поколения автолюбителей по всему миру."
    )

    culture_items = [
        "Фильм «2 Fast 2 Furious» (2003) — серебристый R34 GT-R Брайана О'Коннера",
        "Аниме «Initial D» — R32 GT-R персонажа Сигэхиро Кунимото",
        "Видеоигра «Gran Turismo» — Skyline присутствует с первой части (1997)",
        "Уличные гонки Wangan и Touge на Японских автобанах",
        "JDM-культура: тюнинг, дрифт, стенс — Skyline в центре каждого направления",
    ]
    for item in culture_items:
        doc.add_paragraph(item, style="List Bullet")

    add_highlight_box(
        doc,
        "🏁 «Легенда говорит: на треке нет заменителя GT-R. "
        "На улице — тоже.» — Skyline стал символом японского инженерного гения "
        "и неукротимого духа гонок."
    )

    # --- Заключение ---
    doc.add_heading("Заключение", level=1)
    add_horizontal_line(doc)

    doc.add_paragraph(
        "За более чем 65-летнюю историю Nissan Skyline превратился из обычного седана "
        "в автомобильную икону мирового масштаба. От Hakosuka до R35 GT-R — каждое "
        "поколение вносило свой вклад в автомобильную историю. Сегодня Skyline "
        "и GT-R остаются объектом восхищения коллекционеров, тюнеров и гонщиков, "
        "а цены на классические R32–R34 неуклонно растут, подтверждая статус "
        "культового автомобиля."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— Конец документа —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.italic = True

    doc.save(OUTPUT_FILE)
    print(f"Документ сохранён: {OUTPUT_FILE}")


def main():
    print("Skachivanie izobrazheniy...")
    image_paths = {}
    for key, url in IMAGES.items():
        ext = "jpg"
        path = download_image(url, f"{key}.{ext}", key=key, force=False)
        if path:
            image_paths[key] = path
            print(f"  OK: {key}")
        else:
            print(f"  FAIL: {key}")

    print("\nСоздание документа...")
    build_document(image_paths)
    print("Готово!")


if __name__ == "__main__":
    main()
